"""
outputs/story_generator.py — User Story Üretim Modülü
Sorumlu: Halise İncir

Açıklama:
Analiz edilen fonksiyonel gereksinimleri çevik (Agile) geliştirme süreçlerine
uygun "User Story" formatına dönüştürür. İsteğe bağlı olarak user_stories.docx
dosyasına da export edilebilir.

Çıktı formatı:
    [
        {
            "req_id":              "REQ_001",
            "role":                "kullanıcı",
            "goal":                "sisteme giriş yapabilmek",
            "benefit":             "kişisel hesabıma erişebilmek",
            "acceptance_criteria": ["...", "..."]
        },
        ...
    ]

Notlar:
- Sadece req_type == "FUNCTIONAL" olan gereksinimler işlenir (token tasarrufu).
- LLM çağrısı başarısız olursa ilgili gereksinim için fallback dict döner;
  pipeline çökmez (graceful degradation).
- python-docx eksikse _export_to_docx() ImportError fırlatır (raporda uyarı verilir).
"""

from __future__ import annotations

import concurrent.futures
from pathlib import Path
from typing import List, Optional

from loguru import logger

from core.models import AnalysisReport, Requirement
from modules.llm_client import LLMClient, LLMClientError
from modules.llm_response_utils import extract_json_object
from modules.story_prompts import (
    build_story_generation_system_prompt,
    build_story_generation_user_prompt,
    build_story_generation_batch_system_prompt,
    build_story_generation_batch_user_prompt,
)

# Varsayılan çıktı dizini
_GENERATED_DIR = Path(__file__).parent / "generated"

_log = logger.bind(module="story_generator")


def _make_fallback_story(req: Requirement) -> dict:
    """LLM çağrısı başarısız olduğunda kullanılan temel user story sözlüğü.

    Args:
        req: Dönüştürülemeyen Requirement nesnesi.

    Returns:
        dict: Temel alanları doldurulmuş user story.
    """
    role = req.actors[0] if req.actors else "kullanıcı"
    return {
        "req_id": req.id,
        "role": role,
        "goal": req.text.strip(),
        "benefit": "(LLM analizi mevcut değil)",
        "acceptance_criteria": [f"{req.text.strip()} işlevi çalışır durumda olmalı."],
    }


def _normalize_story(req_id: str, raw: dict) -> dict:
    """LLM'den dönen ham dict'i sözleşmeye uygun hale getirir.

    Eksik alanlar varsayılan değerlerle doldurulur; tip uyumsuzlukları giderilir.

    Args:
        req_id: Gereksinim kimliği (REQ_NNN).
        raw: LLM'den gelen ham sözlük.

    Returns:
        dict: role, goal, benefit, acceptance_criteria, req_id alanları dolu story.
    """
    role = raw.get("role") or "kullanıcı"
    goal = raw.get("goal") or "(hedef belirtilmedi)"
    benefit = raw.get("benefit") or "(fayda belirtilmedi)"

    ac_raw = raw.get("acceptance_criteria", [])
    if isinstance(ac_raw, list):
        acceptance_criteria: List[str] = [str(item) for item in ac_raw if item]
    else:
        acceptance_criteria = [str(ac_raw)] if ac_raw else []

    if not acceptance_criteria:
        acceptance_criteria = ["İlgili işlev hatasız çalışmalı."]

    return {
        "req_id": req_id,
        "role": role,
        "goal": goal,
        "benefit": benefit,
        "acceptance_criteria": acceptance_criteria,
    }


_STORY_CHUNK_SIZE = 15
_STORY_BATCH_MAX_TOKENS = 3072


def _chunk_list(lst: List, size: int) -> List[List]:
    return [lst[i:i + size] for i in range(0, len(lst), size)]


class StoryGenerator:
    """Fonksiyonel gereksinimlerden Agile User Story'leri üretir.

    LLM (Gemini) kullanarak "As a [role], I want [action] so that [benefit]"
    formatında user story'ler üretir. Sadece FUNCTIONAL gereksinimleri işler.
    Gereksinimler chunk'lara bölünerek paralel LLM çağrılarıyla işlenir.

    Args:
        llm_client: Dışarıdan enjekte edilebilir LLM istemcisi (test/DI için).
                    Verilmezse generate() çağrısında otomatik oluşturulur.
    """

    def __init__(self, llm_client: Optional[LLMClient] = None) -> None:
        self._llm_client = llm_client

    def _get_client(self) -> LLMClient:
        if self._llm_client is None:
            self._llm_client = LLMClient(max_output_tokens=_STORY_BATCH_MAX_TOKENS)
        return self._llm_client

    def _generate_single_story(self, req: Requirement) -> dict:
        """Tek bir gereksinim için LLM'e çağrı yapar ve normalize edilmiş story döner.

        Args:
            req: User story'e dönüştürülecek Requirement.

        Returns:
            dict: Normalize edilmiş user story sözlüğü.
        """
        system_prompt = build_story_generation_system_prompt()
        user_prompt = build_story_generation_user_prompt(req)

        client = self._get_client()

        try:
            response = client.chat(
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                metadata={
                    "module": "story_generator",
                    "action": "generate_user_story",
                    "req_id": req.id,
                },
            )
        except LLMClientError:
            _log.exception("LLM çağrısı başarısız (user story üretimi) | req_id={}", req.id)
            raise

        try:
            raw = extract_json_object(response.content)
        except ValueError as exc:
            _log.error(
                "LLM çıktısı JSON olarak çözülemedi | req_id={} preview={}",
                req.id,
                response.content[:300] if response.content else "",
            )
            raise LLMClientError(f"User story çıktısı işlenemedi: {exc}") from exc

        return _normalize_story(req.id, raw)

    def _process_chunk(self, chunk: List[Requirement]) -> List[dict]:
        """Tek bir chunk için LLM çağrısı yapar ve normalize edilmiş story listesi döner."""
        system_prompt = build_story_generation_batch_system_prompt()
        user_prompt = build_story_generation_batch_user_prompt(chunk)
        client = self._get_client()

        response = client.chat(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            metadata={"module": "story_generator", "action": "batch_generate_user_stories"},
        )

        raw = extract_json_object(response.content)
        items = raw if isinstance(raw, list) else (list(raw.values())[0] if isinstance(raw, dict) and raw else [])

        req_index = {r.id: r for r in chunk}
        stories: List[dict] = []
        returned_ids: set = set()

        for item in items:
            if not isinstance(item, dict):
                continue
            req_id = item.get("req_id", "")
            returned_ids.add(req_id)
            stories.append(_normalize_story(req_id, item))

        for req in chunk:
            if req.id not in returned_ids:
                _log.warning("Chunk yanıtında req_id eksik, fallback | req_id={}", req.id)
                stories.append(_make_fallback_story(req_index[req.id]))

        return stories

    def _generate_batch_stories(self, reqs: List[Requirement]) -> List[dict]:
        """Gereksinimleri chunk'lara bölerek paralel LLM çağrılarıyla story'lere dönüştürür.

        Args:
            reqs: Dönüştürülecek Requirement listesi.

        Returns:
            list[dict]: Normalize edilmiş user story listesi (giriş sırasıyla).
        """
        chunks = _chunk_list(reqs, _STORY_CHUNK_SIZE)

        if len(chunks) == 1:
            return self._process_chunk(chunks[0])

        _log.info("Story üretimi {} chunk'a bölündü, paralel işleniyor", len(chunks))
        with concurrent.futures.ThreadPoolExecutor(max_workers=len(chunks)) as executor:
            futures = [executor.submit(self._process_chunk, chunk) for chunk in chunks]
            results: List[dict] = []
            for future in futures:
                try:
                    results.extend(future.result())
                except (LLMClientError, ValueError) as exc:
                    chunk_idx = futures.index(future)
                    _log.warning("Chunk {} başarısız, fallback | hata={}", chunk_idx, exc)
                    results.extend(_make_fallback_story(req) for req in chunks[chunk_idx])
        return results

    def generate(self, report: AnalysisReport) -> List[dict]:
        """AnalysisReport içindeki FUNCTIONAL gereksinimleri User Story'lere dönüştürür.

        Tüm gereksinimler tek bir toplu LLM çağrısıyla işlenir.
        LLM hatası durumunda tüm gereksinimler için fallback dict'ler döner.

        Args:
            report: NLP ve LLM analizleri tamamlanmış rapor nesnesi.

        Returns:
            list[dict]: Her biri {'req_id', 'role', 'goal', 'benefit',
                        'acceptance_criteria'} içeren user story listesi.
                        Yalnızca FUNCTIONAL gereksinimler dahil edilir.
        """
        functional_reqs = [
            req
            for req in report.parsed_doc.requirements
            if req.req_type == "FUNCTIONAL"
        ]

        if not functional_reqs:
            _log.info(
                "User story üretimi atlandı: FUNCTIONAL gereksinim yok "
                "| toplam_gereksinim={}",
                len(report.parsed_doc.requirements),
            )
            return []

        _log.info(
            "User story toplu üretimi başlatıldı | functional_count={}",
            len(functional_reqs),
        )

        try:
            stories = self._generate_batch_stories(functional_reqs)
        except (LLMClientError, ValueError) as exc:
            _log.warning(
                "Toplu user story üretimi başarısız, tüm req'ler için fallback | hata={}",
                exc,
            )
            stories = [_make_fallback_story(req) for req in functional_reqs]

        _log.info("User story üretimi tamamlandı | stories={}", len(stories))
        return stories

    def _export_to_docx(self, stories: List[dict], output_path: Path) -> Path:
        """User Story listesini DOCX formatında dışa aktarır.
        
        Her story için 'Rol, Hedef, Fayda' formatında Türkçe başlıklar ve 
        kabul kriterleri içeren bir Word belgesi oluşturur.

        Args:
            stories: generate() metodundan dönen user story sözlük listesi.
            output_path: Yazılacak .docx dosyasının tam yolu.

        Returns:
            Path: Yazılan dosyanın mutlak yolu.

        Raises:
            ImportError: python-docx paketi kurulu değilse.
        """
        try:
            from docx import Document  # type: ignore
            from docx.shared import Pt  # type: ignore
        except ImportError as exc:
            raise ImportError(
                "python-docx paketi gerekli: pip install python-docx"
            ) from exc

        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading("AutoReq — Kullanıcı Hikayeleri (User Stories)", 0)
        doc.add_paragraph(
            f"Toplam {len(stories)} kullanıcı hikayesi oluşturuldu. "
            "Yalnızca FONKSİYONEL gereksinimler dahil edilmiştir."
        )
        doc.add_paragraph("")

        for story in stories:
            req_id = story.get("req_id", "REQ_??")
            role = story.get("role", "kullanıcı")
            goal = story.get("goal", "(hedef belirtilmedi)")
            benefit = story.get("benefit", "(fayda belirtilmedi)")
            acceptance_criteria: List[str] = story.get("acceptance_criteria") or []

            # Story başlığı
            heading = doc.add_heading(level=2)
            heading.add_run(f"[{req_id}] ").bold = True
            heading.add_run(f"Bir {role} olarak, {benefit} amacıyla {goal} istiyorum.")

            # Acceptance Criteria
            if acceptance_criteria:
                doc.add_paragraph("Kabul Kriterleri:", style="Intense Quote")
                for criterion in acceptance_criteria:
                    doc.add_paragraph(criterion, style="List Bullet")

            doc.add_paragraph("")

        doc.save(str(output_path))
        _log.info("User Stories DOCX dosyası yazıldı | path={} stories={}", output_path, len(stories))
        return output_path

    def generate_and_export(
        self,
        report: AnalysisReport,
        output_path: Optional[Path] = None,
    ) -> tuple[List[dict], Path]:
        """Gereksinimleri user story'lere dönüştürür ve DOCX'e dışa aktarır.

        generate() ve _export_to_docx() metodlarını birleştiren kolaylaştırıcı metod.
        Çıktı path'i belirtilmezse outputs/generated/user_stories.docx kullanılır.

        Args:
            report: NLP ve LLM analizleri tamamlanmış rapor nesnesi.
            output_path: Kaydedilecek .docx dosya yolu. None ise varsayılan kullanılır.

        Returns:
            tuple[list[dict], Path]: (user_story_listesi, docx_dosya_yolu) çifti.
        """
        if output_path is None:
            _GENERATED_DIR.mkdir(parents=True, exist_ok=True)
            output_path = _GENERATED_DIR / "user_stories.docx"

        stories = self.generate(report)
        docx_path = self._export_to_docx(stories, output_path)
        return stories, docx_path
