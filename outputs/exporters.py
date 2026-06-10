"""
outputs/exporters.py — Çoklu Format Dışa Aktarım Modülü
Sorumlu: Halise İncir

Açıklama:
AnalysisReport ve türev verilerini (Backlog, User Stories) farklı formatlarda
dışa aktarır. Desteklenen formatlar:
    - Excel (.xlsx)  → Backlog tablosu (openpyxl)
    - DOCX           → User Stories resmi şablon (python-docx)
    - JSON           → Tam AnalysisReport (Pydantic model_dump_json)

Kullanım:
    from outputs.exporters import export_backlog_xlsx, export_stories_docx, export_report_json

    xlsx_path = export_backlog_xlsx(backlog_items, path)
    docx_path = export_stories_docx(stories, path)
    json_path = export_report_json(report, path)

Notlar:
- openpyxl eksikse export_backlog_xlsx ImportError fırlatır.
- python-docx eksikse export_stories_docx ImportError fırlatır.
- Hedef klasör otomatik oluşturulur (Path.mkdir parents=True, exist_ok=True).
"""

from __future__ import annotations

from pathlib import Path
from typing import List, Optional

from loguru import logger

from core.models import AnalysisReport

_log = logger.bind(module="exporters")

# Varsayılan çıktı dizini
_GENERATED_DIR = Path(__file__).parent / "generated"


# ---------------------------------------------------------------------------
# Excel (XLSX) Dışa Aktarım
# ---------------------------------------------------------------------------


def export_backlog_xlsx(
    backlog: List[dict],
    path: Optional[Path] = None,
    in_memory: bool = False,
) -> Path | io.BytesIO:
    """Product Backlog'u Excel (.xlsx) formatında dışa aktarır.

    Her backlog öğesi için req_id, title, priority_score, story_points,
    type ve depends_on sütunlarını içeren bir tablo oluşturur.

    Args:
        backlog: BacklogGenerator.generate() çıktısı — list[dict].
        path: Yazılacak .xlsx dosyasının yolu. None ise
              outputs/generated/backlog.xlsx kullanılır.
        in_memory: True ise dosyayı disk yerine bellekte (io.BytesIO) üretir.

    Returns:
        Path | io.BytesIO: Yazılan .xlsx dosyasının mutlak yolu veya bellek tamponu.

    Raises:
        ImportError: openpyxl paketi kurulu değilse.
    """
    try:
        import openpyxl  # type: ignore
        from openpyxl.styles import Alignment, Font, PatternFill  # type: ignore
        from openpyxl.utils import get_column_letter  # type: ignore
    except ImportError as exc:
        raise ImportError("openpyxl paketi gerekli: pip install openpyxl") from exc

    import io

    if not in_memory:
        if path is None:
            _GENERATED_DIR.mkdir(parents=True, exist_ok=True)
            path = _GENERATED_DIR / "backlog.xlsx"
        else:
            path.parent.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Product Backlog"

    # Başlık satırı
    headers = ["ID", "Gereksinim", "Öncelik Skoru", "Story Points", "Tür", "Bağımlılıklar"]
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2F4F8F", end_color="2F4F8F", fill_type="solid")

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # Veri satırları
    for row_idx, item in enumerate(backlog, start=2):
        depends_on = item.get("depends_on") or []
        depends_str = ", ".join(str(d) for d in depends_on) if depends_on else "-"

        row_data = [
            item.get("req_id", ""),
            item.get("title", ""),
            item.get("priority_score", 0.0),
            item.get("story_points", 1),
            item.get("type", ""),
            depends_str,
        ]

        # Çift satır arka planı
        row_fill = PatternFill(
            start_color="EEF2FF" if row_idx % 2 == 0 else "FFFFFF",
            end_color="EEF2FF" if row_idx % 2 == 0 else "FFFFFF",
            fill_type="solid",
        )

        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.fill = row_fill
            if col_idx in (3, 4):  # Sayısal sütunlar ortala
                cell.alignment = Alignment(horizontal="center")

    # Sütun genişliklerini otomatik ayarla
    col_widths = [12, 60, 16, 16, 18, 30]
    for col_idx, width in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    # Başlık satırı yüksekliği
    ws.row_dimensions[1].height = 20

    if in_memory:
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        _log.info("Backlog Excel dosyası bellekte (BytesIO) oluşturuldu.")
        return buf

    wb.save(str(path))
    _log.info(
        "Backlog Excel dosyası yazıldı | path={} öğe_sayısı={}", path, len(backlog)
    )
    return path


# ---------------------------------------------------------------------------
# DOCX Dışa Aktarım — User Stories
# ---------------------------------------------------------------------------


def export_stories_docx(
    stories: List[dict],
    path: Optional[Path] = None,
    in_memory: bool = False,
) -> Path | io.BytesIO:
    """User Stories listesini DOCX formatında dışa aktarır.

    Her story için resmi şablon: başlık + 'As a [role], I want [goal] so
    that [benefit].' cümlesi + Acceptance Criteria tablosu.

    Args:
        stories: StoryGenerator.generate() çıktısı — list[dict].
        path: Yazılacak .docx dosyasının yolu. None ise
              outputs/generated/user_stories.docx kullanılır.
        in_memory: True ise dosyayı disk yerine bellekte (io.BytesIO) üretir.

    Returns:
        Path | io.BytesIO: Yazılan .docx dosyasının mutlak yolu veya bellek tamponu.

    Raises:
        ImportError: python-docx paketi kurulu değilse.
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "python-docx paketi gerekli: pip install python-docx"
        ) from exc

    import io

    if not in_memory:
        if path is None:
            _GENERATED_DIR.mkdir(parents=True, exist_ok=True)
            path = _GENERATED_DIR / "user_stories.docx"
        else:
            path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Belge başlığı
    title_para = doc.add_heading("AutoReq — User Stories", 0)
    summary_para = doc.add_paragraph(
        f"Toplam {len(stories)} user story oluşturuldu. "
        "Yalnızca FUNCTIONAL gereksinimler dahil edilmiştir."
    )
    doc.add_paragraph("")

    for story in stories:
        req_id = story.get("req_id", "REQ_??")
        role = story.get("role", "kullanıcı")
        goal = story.get("goal", "(hedef belirtilmedi)")
        benefit = story.get("benefit", "(fayda belirtilmedi)")
        acceptance_criteria: List[str] = story.get("acceptance_criteria") or []

        # Story başlığı
        heading = doc.add_heading(level=2)
        heading.add_run(f"[{req_id}] ").bold = True
        heading.add_run(f"Bir {role} olarak, {benefit} amacıyla {goal} istiyorum.")

        # Acceptance Criteria bölümü
        if acceptance_criteria:
            ac_label = doc.add_paragraph()
            ac_run = ac_label.add_run("Kabul Kriterleri:")
            ac_run.bold = True
            for criterion in acceptance_criteria:
                doc.add_paragraph(criterion, style="List Bullet")

        doc.add_paragraph("")

    if in_memory:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        _log.info("User Stories DOCX dosyası bellekte (BytesIO) oluşturuldu.")
        return buf

    doc.save(str(path))
    _log.info(
        "User Stories DOCX dosyası yazıldı | path={} stories={}",
        path,
        len(stories),
    )
    return path


# ---------------------------------------------------------------------------
# JSON Dışa Aktarım — Tam AnalysisReport
# ---------------------------------------------------------------------------


def export_report_json(
    report: AnalysisReport,
    path: Optional[Path] = None,
    in_memory: bool = False,
) -> Path | io.BytesIO:
    """Tam AnalysisReport'u JSON formatında dışa aktarır.

    Pydantic model_dump_json() kullanılarak tip-güvenli serileştirme yapılır.
    Türkçe karakterler (ş, ğ, İ) bozulmadan korunur (ensure_ascii=False).

    Args:
        report: Serileştirilecek AnalysisReport nesnesi.
        path: Yazılacak .json dosyasının yolu. None ise
              outputs/generated/analysis_report.json kullanılır.
        in_memory: True ise dosyayı disk yerine bellekte (io.BytesIO) üretir.

    Returns:
        Path | io.BytesIO: Yazılan .json dosyasının mutlak yolu veya bellek tamponu.
    """
    import io

    if not in_memory:
        if path is None:
            _GENERATED_DIR.mkdir(parents=True, exist_ok=True)
            path = _GENERATED_DIR / "analysis_report.json"
        else:
            path.parent.mkdir(parents=True, exist_ok=True)

    # Pydantic v2 model_dump_json (ensure_ascii=False için özel serileştirme)
    json_str = report.model_dump_json(indent=2)

    if in_memory:
        buf = io.BytesIO(json_str.encode("utf-8"))
        buf.seek(0)
        _log.info("AnalysisReport JSON dosyası bellekte (BytesIO) oluşturuldu.")
        return buf

    # Türkçe karakterlerin bozulmaması için doğrudan UTF-8 ile yaz
    path.write_text(json_str, encoding="utf-8")

    _log.info(
        "AnalysisReport JSON dosyası yazıldı | path={} boyut_bytes={}",
        path,
        len(json_str.encode("utf-8")),
    )
    return path
